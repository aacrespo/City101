#!/usr/bin/env python3
"""
Convert ARCHITECTURE_DESIGN_DOC.md to a professional DOCX using python-docx.
Handles: headings, paragraphs, tables, code blocks, lists, and ASCII diagrams.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Colors
GOLD = RGBColor(201, 168, 76)      # #c9a84c
DARK_HEADER = RGBColor(26, 26, 46) # #1a1a2e
BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = RGBColor(240, 240, 245)

def set_cell_background(cell, fill):
    """Set table cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def read_markdown(filepath):
    """Read markdown file."""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def extract_sections(content):
    """Extract all H2/H3 sections for TOC."""
    lines = content.split('\n')
    sections = []
    for line in lines:
        if line.startswith('## '):
            sections.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            sections.append(('h3', line[4:].strip()))
    return sections

def add_footer_with_page_numbers(section):
    """Add page numbers to footer."""
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = ''
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def parse_markdown_to_docx(md_content):
    """Parse markdown and build DOCX document."""
    doc = Document()

    # Set A4 page size and margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Add footer with page numbers
    add_footer_with_page_numbers(section)

    # Extract sections for TOC
    sections = extract_sections(md_content)

    # Add title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('City101 Relay-Lock Configurator')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_HEADER

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('Architecture Design Document')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = GOLD

    doc.add_paragraph()  # Spacing

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run('EPFL BA6 Architecture Studio — "Still on the Line"\nTeam: Andrea, Henna, Claude\nMidterm: March 30, 2026')
    meta_run.font.size = Pt(11)
    meta_run.font.italic = True

    doc.add_page_break()

    # Add TOC
    toc_heading = doc.add_paragraph()
    toc_run = toc_heading.add_run('Table of Contents')
    toc_run.font.size = Pt(16)
    toc_run.font.bold = True
    toc_run.font.color.rgb = DARK_HEADER

    for h_type, title in sections:
        indent = 0.5 if h_type == 'h3' else 0
        toc_item = doc.add_paragraph(title, style='List Bullet')
        toc_item.paragraph_format.left_indent = Inches(indent)
        toc_item.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # Parse content
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Skip title (H1)
        if line.startswith('# '):
            i += 1
            continue

        # Skip metadata lines at start
        if i < 10 and (line.startswith('**EPFL') or line.startswith('**Team:') or
                       line.startswith('**Midterm:') or line.strip() == '---'):
            i += 1
            continue

        # Code blocks (triple backtick)
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                # Peek ahead to get language
                lang = line.strip()[3:] if len(line.strip()) > 3 else ''
            else:
                in_code_block = False
                # Add code block
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_para.style = 'Normal'
                    code_para.paragraph_format.left_indent = Inches(0.5)
                    code_para.paragraph_format.space_before = Pt(6)
                    code_para.paragraph_format.space_after = Pt(6)

                    for code_line in code_lines:
                        run = code_para.add_run(code_line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = BLACK

                    # Set background
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'f0f0f5')
                    code_para._element.get_or_add_pPr().append(shading_elm)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables (pipe-delimited)
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue

        if in_table and ('|' not in line or not line.strip().startswith('|')):
            in_table = False
            # Process table
            if table_lines:
                add_table_to_doc(doc, table_lines)
            table_lines = []

        # Headings
        if line.startswith('## '):
            para = doc.add_paragraph(line[3:].strip())
            para.style = 'Heading 2'
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.color.rgb = DARK_HEADER
                run.font.size = Pt(14)
                run.font.bold = True
            i += 1
            continue

        if line.startswith('### '):
            para = doc.add_paragraph(line[4:].strip())
            para.style = 'Heading 3'
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(4)
            for run in para.runs:
                run.font.color.rgb = DARK_HEADER
                run.font.size = Pt(12)
                run.font.bold = True
            i += 1
            continue

        # Mermaid diagrams
        if line.strip().startswith('```mermaid'):
            # Find end of block
            j = i + 1
            while j < len(lines) and not lines[j].strip().startswith('```'):
                j += 1
            # Add placeholder
            diag_para = doc.add_paragraph()
            diag_run = diag_para.add_run('[Diagram: See markdown source]')
            diag_run.italic = True
            diag_run.font.color.rgb = RGBColor(100, 100, 100)
            diag_para.paragraph_format.space_before = Pt(6)
            diag_para.paragraph_format.space_after = Pt(6)
            i = j + 1
            continue

        # Lists (unordered)
        if line.strip().startswith('- '):
            para = doc.add_paragraph(line.strip()[2:], style='List Bullet')
            para.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Lists (ordered)
        if re.match(r'^\d+\.\s', line.strip()):
            match = re.match(r'^(\d+)\.\s(.*)', line.strip())
            if match:
                para = doc.add_paragraph(match.group(2), style='List Number')
                para.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Bold/italic inline formatting
        if line.strip():
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)

            # Process inline markdown
            text = line.strip()
            # Replace **text** with bold, *text* with italic, `code` with code
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)

            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    if part:
                        para.add_run(part)

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        add_table_to_doc(doc, table_lines)

    return doc

def add_table_to_doc(doc, table_lines):
    """Add a markdown table to the document."""
    if not table_lines or len(table_lines) < 2:
        return

    # Parse header row
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]

    # Skip separator line (line 1)
    # Parse data rows (from line 2 onwards)
    data_rows = []
    for line in table_lines[2:]:
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                data_rows.append(cells)

    if not data_rows:
        return

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_background(header_cells[i], 'c9a84c')
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = cell_text

    # Add spacing after table
    doc.add_paragraph()

def main():
    """Main conversion function."""
    print("Reading markdown file...")
    md_content = read_markdown('/sessions/kind-tender-carson/mnt/app_architecture/ARCHITECTURE_DESIGN_DOC.md')

    print("Parsing and building DOCX...")
    doc = parse_markdown_to_docx(md_content)

    output_path = '/sessions/kind-tender-carson/mnt/app_architecture/ARCHITECTURE_DESIGN_DOC.docx'
    print(f"Saving to {output_path}...")
    doc.save(output_path)

    # Verify the file
    print("Verifying output file...")
    try:
        verify_doc = Document(output_path)
        para_count = len(verify_doc.paragraphs)
        table_count = len(verify_doc.tables)
        print(f"✓ File is valid DOCX")
        print(f"  - {para_count} paragraphs")
        print(f"  - {table_count} tables")
        print(f"✓ Output written to: {output_path}")
    except Exception as e:
        print(f"✗ Error verifying file: {e}")
        return False

    return True

if __name__ == '__main__':
    success = main()
    exit(0 if success else 1)
